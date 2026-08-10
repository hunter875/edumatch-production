from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "cv"


def set_font(run, size=8.6, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def bottom_border(p):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def text(p, value, size=8.6, bold=False, color=None):
    r = p.add_run(value)
    set_font(r, size=size, bold=bold, color=color)
    return r


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    text(p, title.upper(), size=10.3, bold=True)
    bottom_border(p)


def bullet(cell, value):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    text(p, value, size=8.2)


def setup():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.32)
    sec.bottom_margin = Inches(0.32)
    sec.left_margin = Inches(0.42)
    sec.right_margin = Inches(0.42)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8.6)
    return doc


def header(doc, vi=True):
    table = doc.add_table(rows=1, cols=2)
    no_borders(table)
    table.autofit = False
    pic, info = table.rows[0].cells
    cell_width(pic, 1.05)
    cell_width(info, 6.2)
    shade(pic, "EDEDED")
    pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = pic.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(27)
    p.paragraph_format.space_after = Pt(27)
    text(p, "ẢNH\n3x4" if vi else "PHOTO\n3x4", size=9, bold=True, color=RGBColor(100, 100, 100))

    p = info.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    text(p, "[HỌ VÀ TÊN]" if vi else "[YOUR NAME]", size=16, bold=True)
    p = info.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    text(p, "Junior Full-Stack / Backend Developer", size=10.5)

    rows = [
        ("Ngày sinh:", "[DD/MM/YYYY]") if vi else ("Date of birth:", "[DD/MM/YYYY]"),
        ("Số điện thoại:", "[số điện thoại]") if vi else ("Phone:", "[phone number]"),
        ("Email:", "[email]"),
        ("Website:", "[GitHub / LinkedIn / Upwork]"),
        ("Địa chỉ:", "TP. Hồ Chí Minh, Việt Nam") if vi else ("Address:", "Ho Chi Minh City, Vietnam"),
    ]
    for label, value in rows:
        p = info.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        text(p, label + " ", size=8.7, bold=True)
        text(p, value, size=8.7)


def entry(doc, date, title, role, items):
    table = doc.add_table(rows=1, cols=2)
    no_borders(table)
    table.autofit = False
    c1, c2 = table.rows[0].cells
    cell_width(c1, 1.25)
    cell_width(c2, 6.05)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = c1.paragraphs[0]
    text(p, date, size=8.4)
    p = c2.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    text(p, title, size=8.9, bold=True)
    if role:
        p = c2.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        text(p, role, size=8.5, bold=True)
    for item in items:
        bullet(c2, item)


def skill_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    no_borders(table)
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cell_width(cells[0], 1.25)
        cell_width(cells[1], 6.05)
        p = cells[0].paragraphs[0]
        text(p, label, size=8.4, bold=True)
        p = cells[1].paragraphs[0]
        text(p, value, size=8.2)


def build_vi():
    doc = setup()
    header(doc, vi=True)

    section(doc, "Mục tiêu nghề nghiệp")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    text(p, "Lập trình viên junior định hướng Full-Stack/Backend, mong muốn tham gia các dự án thực tế về API, database, Docker/cloud deployment và kiểm thử luồng nghiệp vụ.", size=8.4)

    section(doc, "Học vấn")
    entry(doc, "[Năm] - [Năm]", "[Tên trường]", "Chuyên ngành: [Ngành học]", [
        "Học tập và thực hành về lập trình web, cơ sở dữ liệu, API backend và triển khai hệ thống."
    ])

    section(doc, "Dự án / Kinh nghiệm thực hành")
    entry(doc, "2026", "EduMatch - Nền tảng tìm kiếm và quản lý học bổng", "Full-Stack / Backend Developer", [
        "Xây dựng nền tảng web cho học sinh tìm học bổng, nộp đơn và trao đổi với nhà cung cấp.",
        "Phát triển các chức năng: đăng nhập JWT, quản lý học bổng, đơn ứng tuyển, bookmark, chat, thông báo, dashboard provider và trang quản trị.",
        "Xây dựng matching service gợi ý học bổng theo GPA, bậc học, kỹ năng, hình thức học và địa điểm.",
        "Tối ưu một số truy vấn MySQL bằng index và EXPLAIN ANALYZE; đóng gói Docker và deploy thử lên Azure Container Apps bằng GitHub Actions.",
    ])
    entry(doc, "2026", "Doc Automation Engine - Bóc tách Báo cáo PDF và Xuất Word", "Backend Developer", [
        "Xây dựng backend FastAPI cho upload PDF, tạo job bóc tách dữ liệu, kiểm tra kết quả và xuất file Word.",
        "Dùng pdfplumber và regex để trích xuất các trường chính từ báo cáo PCCC theo cấu trúc có sẵn.",
        "Tách tác vụ xử lý nặng sang Celery worker với Redis queue; lưu kết quả bằng PostgreSQL JSONB và lưu file/template bằng MinIO.",
        "Xuất báo cáo Word từ dữ liệu đã tổng hợp bằng docxtpl/Jinja2.",
    ])

    section(doc, "Kỹ năng")
    skill_table(doc, [
        ("Frontend", "React, Next.js, TypeScript, Tailwind CSS, React Query."),
        ("Backend", "Spring Boot, FastAPI, REST API, JWT authentication, role-based access control."),
        ("Database", "MySQL, PostgreSQL, SQL indexing, EXPLAIN ANALYZE, JSONB."),
        ("DevOps", "Docker, Docker Compose, Nginx, GitHub Actions, Azure Container Apps."),
        ("Khác", "Redis, RabbitMQ, Celery, smoke test, API testing, technical documentation."),
    ])

    section(doc, "Hoạt động")
    entry(doc, "2026", "Tự học và thực hành triển khai hệ thống", None, [
        "Thực hành debug full-stack, viết runbook, kiểm thử luồng chính và ghi lại cách vận hành/deploy hệ thống."
    ])
    doc.save(OUT / "CV_ONE_PAGE_VI.docx")


def build_en():
    doc = setup()
    header(doc, vi=False)

    section(doc, "Career Objective")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    text(p, "Junior full-stack/backend developer looking for practical work in API development, database optimization, Docker/cloud deployment, and business-flow testing.", size=8.4)

    section(doc, "Education")
    entry(doc, "[Year] - [Year]", "[Your University]", "Major: [Your Major]", [
        "Studied and practiced web development, databases, backend APIs, and system deployment."
    ])

    section(doc, "Projects / Practical Experience")
    entry(doc, "2026", "EduMatch - Scholarship Search and Management Platform", "Full-Stack / Backend Developer", [
        "Built a web platform where students can search scholarships, submit applications, and communicate with providers.",
        "Developed JWT login, scholarship management, applications, bookmarks, chat, notifications, provider dashboard, and admin pages.",
        "Built a matching service to recommend scholarships based on GPA, education level, skills, study mode, and location.",
        "Improved selected MySQL queries with indexes and EXPLAIN ANALYZE; packaged with Docker and deployed a staging version to Azure Container Apps via GitHub Actions.",
    ])
    entry(doc, "2026", "Doc Automation Engine - PDF Report Extraction and Word Export System", "Backend Developer", [
        "Built a FastAPI backend for uploading PDFs, creating extraction jobs, reviewing extracted data, and exporting Word files.",
        "Used pdfplumber and regex to extract key fields from structured PCCC reports.",
        "Moved heavy processing to Celery workers with Redis queues; stored extracted data in PostgreSQL JSONB and files/templates in MinIO.",
        "Generated Word reports from aggregated data using docxtpl/Jinja2.",
    ])

    section(doc, "Skills")
    skill_table(doc, [
        ("Frontend", "React, Next.js, TypeScript, Tailwind CSS, React Query."),
        ("Backend", "Spring Boot, FastAPI, REST API, JWT authentication, role-based access control."),
        ("Database", "MySQL, PostgreSQL, SQL indexing, EXPLAIN ANALYZE, JSONB."),
        ("DevOps", "Docker, Docker Compose, Nginx, GitHub Actions, Azure Container Apps."),
        ("Other", "Redis, RabbitMQ, Celery, smoke testing, API testing, technical documentation."),
    ])

    section(doc, "Activities")
    entry(doc, "2026", "Self-learning and system deployment practice", None, [
        "Practiced full-stack debugging, runbook writing, main-flow testing, and documenting deployment/operation steps."
    ])
    doc.save(OUT / "CV_ONE_PAGE_EN.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_vi()
    build_en()


if __name__ == "__main__":
    main()
