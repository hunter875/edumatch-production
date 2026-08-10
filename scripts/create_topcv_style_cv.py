from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "cv"


BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(90, 90, 90)
LIGHT_GRAY = "EDEDED"


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def bottom_border(paragraph, color="000000", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title.upper())
    set_font(r, size=12.5, bold=True, color=BLACK)
    bottom_border(p)


def add_text(paragraph, text, size=10, bold=False, color=None):
    r = paragraph.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return r


def add_bullet(cell, text, size=9.6):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.08
    add_text(p, text, size=size)


def add_entry(doc, left, title, subtitle=None, bullets=None, note=None):
    table = doc.add_table(rows=1, cols=2)
    no_table_borders(table)
    table.autofit = False
    left_cell, right_cell = table.rows[0].cells
    set_cell_width(left_cell, 1.45)
    set_cell_width(right_cell, 5.75)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, left, size=9.5)

    p = right_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, title, size=10, bold=True)

    if subtitle:
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_text(p, subtitle, size=9.6, bold=True)

    if bullets:
        for b in bullets:
            add_bullet(right_cell, b)

    if note:
        p = right_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        add_text(p, note, size=9.6, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_skill_row(doc, label, items):
    table = doc.add_table(rows=1, cols=2)
    no_table_borders(table)
    table.autofit = False
    c1, c2 = table.rows[0].cells
    set_cell_width(c1, 1.45)
    set_cell_width(c2, 5.75)
    p = c1.paragraphs[0]
    add_text(p, label, size=9.7, bold=True)
    for item in items:
        add_bullet(c2, item, size=9.6)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.48)
    sec.right_margin = Inches(0.48)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    return doc


def add_header_vi(doc):
    table = doc.add_table(rows=1, cols=2)
    no_table_borders(table)
    table.autofit = False
    img_cell, info_cell = table.rows[0].cells
    set_cell_width(img_cell, 1.45)
    set_cell_width(info_cell, 5.85)

    shade(img_cell, LIGHT_GRAY)
    img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = img_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(42)
    add_text(p, "ẢNH\n3x4", size=12, bold=True, color=GRAY)

    p = info_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "[HỌ VÀ TÊN]", size=18, bold=True)

    p = info_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "Junior Full-Stack / Backend Developer", size=12)

    fields = [
        ("Ngày sinh:", "[DD/MM/YYYY]"),
        ("Giới tính:", "[Nam/Nữ]"),
        ("Số điện thoại:", "[số điện thoại]"),
        ("Email:", "[email]"),
        ("Website:", "[GitHub / LinkedIn / Upwork]"),
        ("Địa chỉ:", "TP. Hồ Chí Minh, Việt Nam"),
    ]
    for label, value in fields:
        p = info_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, label + "  ", size=9.8, bold=True)
        add_text(p, value, size=9.8)


def add_header_en(doc):
    table = doc.add_table(rows=1, cols=2)
    no_table_borders(table)
    table.autofit = False
    img_cell, info_cell = table.rows[0].cells
    set_cell_width(img_cell, 1.45)
    set_cell_width(info_cell, 5.85)

    shade(img_cell, LIGHT_GRAY)
    img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = img_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(42)
    add_text(p, "PHOTO\n3x4", size=12, bold=True, color=GRAY)

    p = info_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "[YOUR NAME]", size=18, bold=True)

    p = info_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "Junior Full-Stack / Backend Developer", size=12)

    fields = [
        ("Date of birth:", "[DD/MM/YYYY]"),
        ("Gender:", "[Male/Female]"),
        ("Phone:", "[phone number]"),
        ("Email:", "[email]"),
        ("Website:", "[GitHub / LinkedIn / Upwork]"),
        ("Address:", "Ho Chi Minh City, Vietnam"),
    ]
    for label, value in fields:
        p = info_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, label + "  ", size=9.8, bold=True)
        add_text(p, value, size=9.8)


def build_vi():
    doc = setup_doc()
    add_header_vi(doc)

    add_section(doc, "Mục tiêu nghề nghiệp")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    add_text(
        p,
        "Tôi là lập trình viên junior định hướng Full-Stack/Backend, mong muốn làm việc trong môi trường thực tế để phát triển kỹ năng xây dựng API, tối ưu database, xử lý bug, triển khai Docker/cloud và viết tài liệu kỹ thuật rõ ràng. Tôi có kinh nghiệm thực hành qua các dự án microservices, hệ thống xử lý tài liệu, Redis/RabbitMQ, CI/CD và Azure Container Apps.",
        size=9.8,
    )

    add_section(doc, "Học vấn")
    add_entry(
        doc,
        "[Năm bắt đầu] - [Năm kết thúc]",
        "[Tên trường]",
        "Chuyên ngành: [Ngành học của bạn]",
        [
            "Môn học/liên quan: Cơ sở dữ liệu, Lập trình Web, Kiến trúc phần mềm, Mạng máy tính.",
            "Định hướng học tập: Backend, API, database, triển khai hệ thống và kiểm thử luồng nghiệp vụ.",
        ],
    )

    add_section(doc, "Dự án / Kinh nghiệm thực hành")
    add_entry(
        doc,
        "2026",
        "EduMatch - Nền tảng tìm kiếm và quản lý học bổng",
        "Full-Stack / Backend Developer",
        [
            "Xây dựng nền tảng web cho phép học sinh tìm kiếm học bổng, nộp đơn và trao đổi với nhà cung cấp.",
            "Phát triển các chức năng chính: đăng nhập JWT, quản lý học bổng, nộp đơn, bookmark, chat, thông báo, dashboard provider và trang quản trị.",
            "Xây dựng matching service để gợi ý học bổng phù hợp dựa trên GPA, bậc học, kỹ năng, hình thức học và địa điểm.",
            "Tối ưu một số truy vấn MySQL cho danh sách học bổng, thống kê đơn ứng tuyển và bookmark bằng index và EXPLAIN ANALYZE.",
            "Đóng gói hệ thống bằng Docker, cấu hình Nginx gateway và triển khai thử lên Azure Container Apps bằng GitHub Actions.",
        ],
    )
    add_entry(
        doc,
        "2026",
        "Doc Automation Engine - Bóc tách Báo cáo PDF và Xuất Word",
        "Backend Developer",
        [
            "Xây dựng backend FastAPI cho luồng upload báo cáo PDF, tạo job bóc tách dữ liệu, kiểm tra kết quả và xuất file Word.",
            "Dùng pdfplumber và regex để trích xuất các trường chính từ báo cáo PCCC theo cấu trúc có sẵn.",
            "Tách các tác vụ xử lý nặng sang Celery worker với Redis queue để API không phải chờ xử lý trực tiếp.",
            "Lưu metadata và kết quả bóc tách bằng PostgreSQL JSONB, lưu file gốc và template bằng MinIO.",
            "Xuất báo cáo Word từ dữ liệu đã tổng hợp bằng docxtpl/Jinja2.",
        ],
    )

    add_section(doc, "Hoạt động")
    add_entry(
        doc,
        "2026",
        "Tự học và thực hành triển khai hệ thống",
        None,
        [
            "Thực hành Docker Compose, Azure Container Apps, GitHub Actions, Redis/RabbitMQ, smoke test và rollback.",
            "Viết tài liệu kỹ thuật để ghi lại cách debug, tối ưu và vận hành hệ thống.",
        ],
    )

    add_section(doc, "Chứng chỉ")
    add_entry(doc, "[Năm]", "[Tên chứng chỉ nếu có]", None, ["Có thể bỏ mục này nếu chưa có chứng chỉ."])

    add_section(doc, "Kỹ năng")
    add_skill_row(
        doc,
        "Kỹ năng cứng",
        [
            "Frontend: React, Next.js, TypeScript, Tailwind CSS, React Query.",
            "Backend: Spring Boot, FastAPI, REST API, JWT authentication, role-based access control.",
            "Database/Infra: MySQL, PostgreSQL, Redis, RabbitMQ, Celery, Docker, Nginx, Azure Container Apps, GitHub Actions.",
        ],
    )
    add_skill_row(
        doc,
        "Kỹ năng mềm",
        [
            "Debug theo luồng: frontend network, gateway, service log, database query, dependency.",
            "Giao tiếp rõ ràng, viết tài liệu dễ hiểu, cẩn thận khi kiểm thử luồng người dùng.",
            "Sẵn sàng học công nghệ mới và nhận việc nhỏ để tích lũy kinh nghiệm thực tế.",
        ],
    )

    add_section(doc, "Người giới thiệu")
    p = doc.add_paragraph()
    add_text(p, "Sẽ cung cấp khi được yêu cầu.", size=9.8)

    add_section(doc, "Sở thích")
    p = doc.add_paragraph()
    add_text(p, "Tìm hiểu backend, hệ thống phân tán, cloud deployment, đọc tài liệu kỹ thuật và xây dựng sản phẩm web.", size=9.8)
    doc.save(OUT / "CV_TOPCV_STYLE_VI.docx")


def build_en():
    doc = setup_doc()
    add_header_en(doc)

    add_section(doc, "Career Objective")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    add_text(
        p,
        "Junior full-stack/backend developer looking for practical software development work involving API development, database optimization, bug fixing, Docker/cloud deployment, and clear technical documentation. I have hands-on project experience with microservices, document processing systems, Redis/RabbitMQ, CI/CD, and Azure Container Apps.",
        size=9.8,
    )

    add_section(doc, "Education")
    add_entry(
        doc,
        "[Start year] - [End year]",
        "[Your University]",
        "Major: [Your Major]",
        [
            "Relevant coursework: Database Systems, Web Development, Software Architecture, Computer Networks.",
            "Learning focus: backend development, API design, database performance, deployment, and business-flow testing.",
        ],
    )

    add_section(doc, "Projects / Practical Experience")
    add_entry(
        doc,
        "2026",
        "EduMatch - Scholarship Search and Management Platform",
        "Full-Stack / Backend Developer",
        [
            "Built a web platform where students can search scholarships, submit applications, and communicate with scholarship providers.",
            "Developed core features including JWT login, scholarship management, applications, bookmarks, chat, notifications, provider dashboard, and admin pages.",
            "Built a matching service to recommend scholarships based on GPA, education level, skills, study mode, and location.",
            "Improved selected MySQL queries for scholarship lists, application statistics, and bookmark lookup using indexes and EXPLAIN ANALYZE.",
            "Packaged the system with Docker, configured an Nginx gateway, and deployed a staging version to Azure Container Apps using GitHub Actions.",
        ],
    )
    add_entry(
        doc,
        "2026",
        "Doc Automation Engine - PDF Report Extraction and Word Export System",
        "Backend Developer",
        [
            "Built a FastAPI backend for uploading PDF reports, creating extraction jobs, reviewing extracted data, and exporting Word files.",
            "Used pdfplumber and regex to extract key fields from structured PCCC reports.",
            "Moved heavy document-processing tasks to Celery workers with Redis queues so the API does not block on long-running jobs.",
            "Stored job metadata and extracted data in PostgreSQL JSONB, and stored uploaded files/templates in MinIO.",
            "Generated Word reports from aggregated data using docxtpl/Jinja2.",
        ],
    )

    add_section(doc, "Activities")
    add_entry(
        doc,
        "2026",
        "Self-learning and system deployment practice",
        None,
        [
            "Practiced Docker Compose, Azure Container Apps, GitHub Actions, Redis/RabbitMQ, smoke testing, and rollback workflows.",
            "Wrote technical notes and runbooks to document debugging, optimization, and system operation steps.",
        ],
    )

    add_section(doc, "Certificates")
    add_entry(doc, "[Year]", "[Certificate name if available]", None, ["This section can be removed if you do not have certificates yet."])

    add_section(doc, "Skills")
    add_skill_row(
        doc,
        "Technical",
        [
            "Frontend: React, Next.js, TypeScript, Tailwind CSS, React Query.",
            "Backend: Spring Boot, FastAPI, REST API, JWT authentication, role-based access control.",
            "Database/Infra: MySQL, PostgreSQL, Redis, RabbitMQ, Celery, Docker, Nginx, Azure Container Apps, GitHub Actions.",
        ],
    )
    add_skill_row(
        doc,
        "Soft skills",
        [
            "Debugging across layers: frontend network, gateway, service logs, database queries, and dependencies.",
            "Clear communication, technical documentation, careful testing of important user flows.",
            "Willing to learn new technologies and start with small practical tasks to build real experience.",
        ],
    )

    add_section(doc, "References")
    p = doc.add_paragraph()
    add_text(p, "Available upon request.", size=9.8)

    add_section(doc, "Interests")
    p = doc.add_paragraph()
    add_text(p, "Backend development, distributed systems, cloud deployment, technical documentation, and building web products.", size=9.8)
    doc.save(OUT / "CV_TOPCV_STYLE_EN.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_vi()
    build_en()


if __name__ == "__main__":
    main()
