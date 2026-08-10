from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "cv"


BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders_none(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="1F4E79", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title.upper(), bold=True, color=BLUE, size=10)
    add_bottom_border(p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    p.add_run(text)
    return p


def normal(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    p.add_run(text)
    return p


def project_header(doc, title, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, title, bold=True, size=10.2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, tech, italic=True, color=GRAY, size=8.8)


def apply_base_style(doc):
    section_obj = doc.sections[0]
    section_obj.top_margin = Inches(0.45)
    section_obj.bottom_margin = Inches(0.45)
    section_obj.left_margin = Inches(0.55)
    section_obj.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(8.8)


def add_header(doc, name, title, contact):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    add_run(p, name, bold=True, color=BLUE, size=19)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    add_run(p, title, bold=True, size=10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    add_run(p, contact, color=GRAY, size=8.6)


def build_english(path: Path):
    doc = Document()
    apply_base_style(doc)
    add_header(
        doc,
        "[YOUR NAME]",
        "Junior Full-Stack / Backend Developer",
        "Ho Chi Minh City, Vietnam | [phone] | [email] | [GitHub] | [LinkedIn/Upwork]",
    )

    section(doc, "Summary")
    normal(
        doc,
        "Junior full-stack developer with hands-on project experience building and debugging web applications across frontend, backend, database, Docker, and cloud deployment. Comfortable with React/Next.js, Spring Boot, FastAPI, MySQL, PostgreSQL, Redis, RabbitMQ, Docker, GitHub Actions, and Azure Container Apps.",
    )

    section(doc, "Technical Skills")
    normal(
        doc,
        "Frontend: React, Next.js, TypeScript, Tailwind CSS, React Query\n"
        "Backend: Spring Boot, FastAPI, REST API, JWT authentication, role-based access control\n"
        "Database/Infra: MySQL, PostgreSQL, Redis, RabbitMQ, Celery, Docker, Nginx, Azure Container Apps, GitHub Actions",
    )

    section(doc, "Projects")
    project_header(
        doc,
        "EduMatch - Microservices Scholarship Matching Platform",
        "React, Next.js, Spring Boot, FastAPI, MySQL, PostgreSQL, Redis, RabbitMQ, Docker, Nginx, Azure Container Apps",
    )
    for item in [
        "Built and debugged student, provider, and admin workflows including scholarship browsing, authentication, applications, chat, notifications, dashboards, and moderation.",
        "Improved matching service design using rule-based scoring, batch scoring, PostgreSQL read models, recommendation cache, and worker-based precomputation.",
        "Optimized database performance with load-test data, EXPLAIN ANALYZE, and indexes for scholarship listing, application counts, and bookmark lookup.",
        "Deployed staging to Azure Container Apps with ACR, Nginx gateway, GitHub Actions CI/CD, smoke tests, rollback workflow, and App Insights observability.",
        "Debugged real cloud issues: CORS registration failure, backend unavailable errors, Redis/RabbitMQ internal TCP routing, old active revisions, cold starts, and scale-down.",
    ]:
        bullet(doc, item)

    project_header(
        doc,
        "Doc Automation Engine - PDF Report Extraction and Word Export System",
        "FastAPI, Celery, Redis, PostgreSQL, SQLAlchemy, MinIO, pdfplumber, docxtpl, Docker Compose, Next.js",
    )
    for item in [
        "Reviewed and worked with a layered FastAPI architecture for API routes, application services, domain models, infrastructure, extraction engines, and worker tasks.",
        "Analyzed a document pipeline using deterministic PDF parsing with pdfplumber and regex, plus optional asynchronous LLM enrichment.",
        "Reviewed Celery worker design with Redis queues for extraction, enrichment, document processing, and scheduled background jobs.",
        "Reviewed PostgreSQL JSONB job models, MinIO/S3-compatible storage, checksum-based deduplication, human review flow, and Word template export.",
        "Identified production-readiness issues: missing frontend API library files, weak dev secrets, wildcard CORS, startup DDL changes, missing CI, N+1 member query, and heavy in-process extraction endpoint.",
    ]:
        bullet(doc, item)

    section(doc, "Selected Engineering Work")
    for item in [
        "Database optimization: seed/load-test data, composite indexes, query plan comparison, and performance documentation.",
        "API design: /api/v1 contracts, pagination rules, error response shape, batch endpoints, idempotency notes, and performance budgets.",
        "Cloud deployment: Azure Container Apps staging, secrets, service logs, smoke tests, rollback notes, and maintenance runbooks.",
    ]:
        bullet(doc, item)

    section(doc, "Education")
    normal(doc, "[Your University] - [Degree / Major] | [Expected Graduation Year or Graduation Year]")
    doc.save(path)


def build_vietnamese(path: Path):
    doc = Document()
    apply_base_style(doc)
    add_header(
        doc,
        "[HỌ VÀ TÊN]",
        "Lập trình viên Full-Stack / Backend Junior",
        "TP. Hồ Chí Minh, Việt Nam | [số điện thoại] | [email] | [GitHub] | [LinkedIn/Upwork]",
    )

    section(doc, "Tóm Tắt")
    normal(
        doc,
        "Lập trình viên junior có kinh nghiệm thực hành qua các dự án web full-stack, bao gồm frontend, backend, database, Docker và deploy cloud. Đã làm việc với React/Next.js, Spring Boot, FastAPI, MySQL, PostgreSQL, Redis, RabbitMQ, Docker, GitHub Actions và Azure Container Apps.",
    )

    section(doc, "Kỹ Năng Kỹ Thuật")
    normal(
        doc,
        "Frontend: React, Next.js, TypeScript, Tailwind CSS, React Query\n"
        "Backend: Spring Boot, FastAPI, REST API, JWT authentication, phân quyền theo role\n"
        "Database/Infra: MySQL, PostgreSQL, Redis, RabbitMQ, Celery, Docker, Nginx, Azure Container Apps, GitHub Actions",
    )

    section(doc, "Dự Án")
    project_header(
        doc,
        "EduMatch - Hệ thống Microservices cho Học bổng và Matching",
        "React, Next.js, Spring Boot, FastAPI, MySQL, PostgreSQL, Redis, RabbitMQ, Docker, Nginx, Azure Container Apps",
    )
    for item in [
        "Phát triển và debug các luồng student, provider và admin: danh sách học bổng, đăng nhập/đăng ký, nộp đơn, chat, notification, dashboard và moderation.",
        "Cải tiến matching service với rule-based scoring, batch scoring, PostgreSQL read model, recommendation cache và worker precompute.",
        "Tối ưu database bằng load-test data, EXPLAIN ANALYZE và index cho scholarship list, application count và bookmark lookup.",
        "Deploy staging lên Azure Container Apps với ACR, Nginx gateway, GitHub Actions CI/CD, smoke test, rollback workflow và App Insights observability.",
        "Debug các lỗi cloud thực tế: CORS khi register, backend unavailable, Redis/RabbitMQ internal TCP routing, old active revisions, cold start và scale-down.",
    ]:
        bullet(doc, item)

    project_header(
        doc,
        "Doc Automation Engine - Bóc tách Báo cáo PDF và Xuất Word",
        "FastAPI, Celery, Redis, PostgreSQL, SQLAlchemy, MinIO, pdfplumber, docxtpl, Docker Compose, Next.js",
    )
    for item in [
        "Review và làm việc với kiến trúc FastAPI gồm API routes, application services, domain models, infrastructure, extraction engines và worker tasks.",
        "Phân tích pipeline tài liệu dùng deterministic PDF parsing bằng pdfplumber và regex, kèm optional LLM enrichment chạy bất đồng bộ.",
        "Review thiết kế Celery worker với Redis queue cho extraction, enrichment, document processing và scheduled background jobs.",
        "Review PostgreSQL JSONB job models, MinIO/S3-compatible storage, checksum deduplication, human review flow và Word template export.",
        "Phát hiện các điểm cần harden: thiếu frontend API library files, default dev secrets, wildcard CORS, startup DDL changes, thiếu CI, N+1 member query và endpoint extraction nặng chạy trong request path.",
    ]:
        bullet(doc, item)

    section(doc, "Công Việc Kỹ Thuật Tiêu Biểu")
    for item in [
        "Tối ưu database: seed/load-test data, composite index, so sánh query plan và viết báo cáo performance.",
        "Thiết kế API: /api/v1 contract, pagination, error response, batch endpoint, idempotency note và performance budget.",
        "Deploy cloud: Azure Container Apps staging, secrets, service logs, smoke test, rollback note và maintenance runbook.",
    ]:
        bullet(doc, item)

    section(doc, "Học Vấn")
    normal(doc, "[Tên trường] - [Ngành học / Bằng cấp] | [Năm tốt nghiệp dự kiến hoặc năm tốt nghiệp]")
    doc.save(path)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_english(OUT_DIR / "CV_FULLSTACK_BACKEND_EN.docx")
    build_vietnamese(OUT_DIR / "CV_FULLSTACK_BACKEND_VI.docx")


if __name__ == "__main__":
    main()
