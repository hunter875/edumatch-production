from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "cv"
BLUE = RGBColor(31, 60, 136)
GRAY = RGBColor(85, 85, 85)
LIGHT = RGBColor(245, 247, 251)


def set_run(run, size=9, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_text(paragraph, value, size=9, bold=False, color=None):
    run = paragraph.add_run(value)
    set_run(run, size=size, bold=bold, color=color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def table_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def bottom_rule(paragraph, color="000000", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.42)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.header_distance = Inches(0.15)
    sec.footer_distance = Inches(0.15)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8.8)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    return doc


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_text(p, title.upper(), size=9.4, bold=True)
    bottom_rule(p)


def header(doc, vi=True):
    table = doc.add_table(rows=1, cols=2)
    no_borders(table)
    table.autofit = False
    left, right = table.rows[0].cells
    table_width(left, 5.0)
    table_width(right, 2.2)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_text(p, "[HỌ VÀ TÊN]" if vi else "[YOUR NAME]", size=18, bold=True, color=BLUE)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "Junior Full-Stack / Backend Developer", size=10.2, bold=True)

    contact = (
        "[số điện thoại]  |  [email]  |  TP. Hồ Chí Minh  |  GitHub / LinkedIn / Upwork"
        if vi
        else "[phone]  |  [email]  |  Ho Chi Minh City  |  GitHub / LinkedIn / Upwork"
    )
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_text(p, contact, size=8.5, color=GRAY)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "Backend-focused\nFull-stack\nDeveloper", size=9.2, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    bottom_rule(p, color="1F3C88", size="10")


def body_paragraph(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, value, size=8.7)


def entry(doc, meta, title, role, bullets):
    table = doc.add_table(rows=1, cols=2)
    no_borders(table)
    table.autofit = False
    c1, c2 = table.rows[0].cells
    table_width(c1, 1.15)
    table_width(c2, 6.05)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = c1.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, meta, size=8.2, color=GRAY)

    p = c2.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, title, size=8.8, bold=True)
    if role:
        add_text(p, f" — {role}", size=8.5, bold=True, color=GRAY)

    for item in bullets:
        p = c2.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.08)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_text(p, item, size=8.35)


def skills(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    no_borders(table)
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        table_width(cells[0], 1.15)
        table_width(cells[1], 6.05)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, label, size=8.4, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, size=8.25)


def build_vi():
    doc = setup_doc()
    header(doc, vi=True)

    section(doc, "Mục tiêu nghề nghiệp")
    body_paragraph(
        doc,
        "Junior developer định hướng Full-Stack/Backend, muốn làm các công việc thực tế về xây dựng API, xử lý dữ liệu, tối ưu truy vấn và triển khai hệ thống bằng Docker/cloud.",
    )

    section(doc, "Học vấn")
    entry(
        doc,
        "[Năm] - [Năm]",
        "[Tên trường]",
        "Chuyên ngành: [Ngành học]",
        ["Thực hành lập trình web, cơ sở dữ liệu, API backend và triển khai hệ thống."],
    )

    section(doc, "Dự án")
    entry(
        doc,
        "2026",
        "EduMatch - Nền tảng tìm kiếm và quản lý học bổng",
        "Full-Stack / Backend Developer",
        [
            "Xây dựng các luồng chính: đăng nhập JWT, danh sách học bổng, nộp đơn, bookmark, chat, notification, dashboard provider và trang admin.",
            "Phát triển matching service gợi ý học bổng theo GPA, bậc học, kỹ năng, hình thức học và địa điểm.",
            "Tối ưu một số truy vấn MySQL bằng index và EXPLAIN ANALYZE; đóng gói Docker, cấu hình Nginx gateway và deploy thử lên Azure Container Apps.",
        ],
    )
    entry(
        doc,
        "2026",
        "Doc Automation Engine - Bóc tách PDF và xuất Word",
        "Backend Developer",
        [
            "Xây dựng backend FastAPI cho upload PDF, tạo job bóc tách dữ liệu, kiểm tra kết quả và xuất file Word.",
            "Dùng pdfplumber/regex để trích xuất trường chính từ báo cáo PCCC; đưa tác vụ nặng sang Celery worker với Redis queue.",
            "Lưu metadata/kết quả bằng PostgreSQL JSONB, lưu file bằng MinIO và xuất Word bằng docxtpl/Jinja2.",
        ],
    )

    section(doc, "Kỹ năng")
    skills(
        doc,
        [
            ("Frontend", "React, Next.js, TypeScript, Tailwind CSS, React Query."),
            ("Backend", "Spring Boot, FastAPI, REST API, JWT, RBAC."),
            ("Database", "MySQL, PostgreSQL, SQL indexing, EXPLAIN ANALYZE, JSONB."),
            ("DevOps", "Docker, Docker Compose, Nginx, GitHub Actions, Azure Container Apps."),
            ("Tools", "Redis, RabbitMQ, Celery, API testing, smoke test, Git."),
        ],
    )

    section(doc, "Thông tin thêm")
    skills(
        doc,
        [
            ("Ngôn ngữ", "Tiếng Việt: bản ngữ. Tiếng Anh: đọc tài liệu kỹ thuật, giao tiếp cơ bản."),
            ("Định hướng", "Backend/API development, database optimization, cloud deployment, automation tools."),
        ],
    )
    doc.save(OUT / "CV_ONE_PAGE_CLEAN_VI.docx")


def build_en():
    doc = setup_doc()
    header(doc, vi=False)

    section(doc, "Career Objective")
    body_paragraph(
        doc,
        "Junior full-stack/backend developer looking for practical work in API development, data processing, query optimization, and Docker/cloud deployment.",
    )

    section(doc, "Education")
    entry(
        doc,
        "[Year] - [Year]",
        "[Your University]",
        "Major: [Your Major]",
        ["Practiced web development, databases, backend APIs, and system deployment."],
    )

    section(doc, "Projects")
    entry(
        doc,
        "2026",
        "EduMatch - Scholarship Search and Management Platform",
        "Full-Stack / Backend Developer",
        [
            "Built main flows: JWT login, scholarship listing, applications, bookmarks, chat, notifications, provider dashboard, and admin pages.",
            "Developed a matching service to recommend scholarships based on GPA, education level, skills, study mode, and location.",
            "Improved selected MySQL queries with indexes and EXPLAIN ANALYZE; packaged with Docker, configured Nginx gateway, and deployed a staging version to Azure Container Apps.",
        ],
    )
    entry(
        doc,
        "2026",
        "Doc Automation Engine - PDF Extraction and Word Export",
        "Backend Developer",
        [
            "Built a FastAPI backend for PDF upload, extraction jobs, result review, and Word export.",
            "Used pdfplumber/regex to extract key fields from PCCC reports; moved heavy tasks to Celery workers with Redis queues.",
            "Stored metadata/results in PostgreSQL JSONB, stored files in MinIO, and generated Word reports with docxtpl/Jinja2.",
        ],
    )

    section(doc, "Skills")
    skills(
        doc,
        [
            ("Frontend", "React, Next.js, TypeScript, Tailwind CSS, React Query."),
            ("Backend", "Spring Boot, FastAPI, REST API, JWT, RBAC."),
            ("Database", "MySQL, PostgreSQL, SQL indexing, EXPLAIN ANALYZE, JSONB."),
            ("DevOps", "Docker, Docker Compose, Nginx, GitHub Actions, Azure Container Apps."),
            ("Tools", "Redis, RabbitMQ, Celery, API testing, smoke testing, Git."),
        ],
    )

    section(doc, "Additional")
    skills(
        doc,
        [
            ("Languages", "Vietnamese: native. English: technical reading, basic communication."),
            ("Focus", "Backend/API development, database optimization, cloud deployment, automation tools."),
        ],
    )
    doc.save(OUT / "CV_ONE_PAGE_CLEAN_EN.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_vi()
    build_en()


if __name__ == "__main__":
    main()
